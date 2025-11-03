import os
import streamlit as st
from dotenv import load_dotenv
import time
from typing import List
import PyPDF2
from io import BytesIO
import psycopg2
from psycopg2.extras import RealDictCursor
from datetime import datetime
import re

load_dotenv()

# Environment setup
os.environ["USER_AGENT"] = os.getenv("USER_AGENT", "ai-content-generator/0.1")
os.environ['STREAMLIT_SERVER_FILE_WATCHER_TYPE'] = 'none'

# API key
GROQ_API_KEY = st.secrets.get("API_KEY") or os.getenv("API_KEY")
if not GROQ_API_KEY:
    raise ValueError("API_KEY not found in secrets or environment variables.")

# Database URL
DATABASE_URL = st.secrets.get("DATABASE_URL") or os.getenv("DATABASE_URL")
if not DATABASE_URL:
    raise ValueError("DATABASE_URL not found in secrets or environment variables.")

# Import the bot
try:
    from bot import MultiUserContentCreator, TWEET_TEMPLATES
    print("Successfully imported MultiUserContentCreator")
except ImportError as e:
    st.error(f"Could not import from bot.py: {e}")
    st.stop()


# ============================================
# USER AUTHENTICATION SYSTEM
# ============================================

class UserAuthSystem:
    def __init__(self, db_url):
        self.db_url = db_url
        self.init_database()
    
    def get_connection(self):
        """Create database connection"""
        return psycopg2.connect(self.db_url)
    
    def init_database(self):
        """Create users table if it doesn't exist"""
        try:
            conn = self.get_connection()
            cur = conn.cursor()
            
            cur.execute('''CREATE TABLE IF NOT EXISTS users
                         (id SERIAL PRIMARY KEY,
                          email VARCHAR(255) UNIQUE NOT NULL,
                          created_at TIMESTAMP NOT NULL,
                          last_login TIMESTAMP NOT NULL,
                          usage_count INTEGER DEFAULT 0)''')
            
            cur.execute('''CREATE INDEX IF NOT EXISTS idx_users_email 
                           ON users(email)''')
            
            conn.commit()
            cur.close()
            conn.close()
        except Exception as e:
            st.error(f"Database initialization error: {e}")
    
    def validate_email(self, email):
        """Simple email validation"""
        pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        return re.match(pattern, email) is not None
    
    def add_user(self, email):
        """Add new user or update last login"""
        try:
            conn = self.get_connection()
            cur = conn.cursor()
            
            try:
                cur.execute('''INSERT INTO users (email, created_at, last_login) 
                               VALUES (%s, %s, %s)''',
                            (email, datetime.now(), datetime.now()))
                conn.commit()
                cur.close()
                conn.close()
                return {"success": True, "new_user": True}
            except psycopg2.IntegrityError:
                conn.rollback()
                cur.execute('''UPDATE users 
                               SET last_login = %s 
                               WHERE email = %s''',
                            (datetime.now(), email))
                conn.commit()
                cur.close()
                conn.close()
                return {"success": True, "new_user": False}
        except Exception as e:
            st.error(f"Error adding user: {e}")
            return {"success": False, "new_user": False}
    
    def increment_usage(self, email):
        """Increment usage count for analytics"""
        try:
            conn = self.get_connection()
            cur = conn.cursor()
            cur.execute('''UPDATE users 
                           SET usage_count = usage_count + 1 
                           WHERE email = %s''', (email,))
            conn.commit()
            cur.close()
            conn.close()
        except Exception as e:
            st.error(f"Error incrementing usage: {e}")
    
    def get_user_stats(self):
        """Get total users and usage stats"""
        try:
            conn = self.get_connection()
            cur = conn.cursor(cursor_factory=RealDictCursor)
            
            cur.execute("SELECT COUNT(*) as total_users FROM users")
            total_users = cur.fetchone()['total_users']
            
            cur.execute("SELECT COALESCE(SUM(usage_count), 0) as total_usage FROM users")
            total_usage = cur.fetchone()['total_usage']
            
            cur.close()
            conn.close()
            return {"total_users": total_users, "total_usage": total_usage}
        except Exception as e:
            st.error(f"Error getting stats: {e}")
            return {"total_users": 0, "total_usage": 0}


def show_email_gate():
    """Show email collection screen"""
    st.title("🧵 AI Twitter Content Creator")
    
    st.markdown("""
    ### Welcome! 👋
    
    Create engaging, research-backed Twitter content for crypto projects using AI.
    
    **It's completely FREE while in beta!** Just enter your email to get started.
    """)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        ### 📝 Templates Available:
        - Educational
        - Promotional
        - Thread
        - Engagement
        - Casual
        - Storytelling
        """)
    
    with col2:
        st.markdown("""
        ### 🎯 Features:
        - Research-backed content
        - Multiple templates
        - Custom brand voice
        - Project management
        - Content tweaking
        """)
    
    with col3:
        st.markdown("""
        ### 📄 Supports:
        - PDF files
        - Word documents
        - Markdown files
        - Website URLs
        - Text paste
        """)
    
    st.divider()
    
    auth_system = UserAuthSystem(DATABASE_URL)
    
    with st.form("email_form"):
        col1, col2 = st.columns([2, 1])
        
        with col1:
            email = st.text_input(
                "📧 Enter your email to get started:",
                placeholder="you@example.com"
            )
        
        with col2:
            st.write("")
            st.write("")
            submit = st.form_submit_button("🚀 Start Creating", use_container_width=True)
        
        if submit:
            if not email:
                st.error("Please enter your email")
            elif not auth_system.validate_email(email):
                st.error("Please enter a valid email address")
            else:
                result = auth_system.add_user(email)
                if result['success']:
                    st.session_state.user_email = email
                    st.session_state.email_provided = True
                    
                    if result['new_user']:
                        st.success(f"🎉 Welcome aboard! You're all set.")
                    else:
                        st.success(f"👋 Welcome back!")
                    
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error("Something went wrong. Please try again.")
    
    st.divider()
    st.caption("💡 Free during beta • No credit card required • Cancel anytime")


# ============================================
# FILE PROCESSING
# ============================================

def extract_file_content(uploaded_file):
    """Extract content from uploaded file"""
    try:
        file_content = uploaded_file.read()
        file_name = uploaded_file.name.lower()
        
        if file_name.endswith('.txt') or file_name.endswith('.md'):
            return file_content.decode('utf-8')
        elif file_name.endswith('.pdf'):
            try:
                pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
                return "\n".join([page.extract_text() for page in pdf_reader.pages])
            except Exception as e:
                st.error(f"Error reading PDF: {str(e)}")
                return None
        elif file_name.endswith(('.docx', '.doc')):
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(BytesIO(file_content))
                return "\n".join([paragraph.text for paragraph in doc.paragraphs])
            except ImportError:
                st.error("python-docx not available")
                return None
            except Exception as e:
                st.error(f"Error reading Word document: {str(e)}")
                return None
        else:
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                st.error(f"Cannot decode file {uploaded_file.name}")
                return None
    except Exception as e:
        st.error(f"Error processing {uploaded_file.name}: {str(e)}")
        return None


# ============================================
# CPU-OPTIMIZED PROJECT DATA PROCESSING
# ============================================

def process_project_data(creator, user_id, project_name, urls_list, uploaded_files, additional_text):
    """CPU-optimized project data processing with progress tracking"""
    try:
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Calculate total steps
        total_steps = 0
        if urls_list: total_steps += 1
        if uploaded_files: total_steps += len(uploaded_files)
        if additional_text: total_steps += 1
        total_steps += 1  # For embedding
        
        current_step = 0
        
        # ═══ STEP 1: Process URLs ═══
        if urls_list:
            status_text.text(f"📥 Loading {len(urls_list)} URLs...")
            status_text.caption("⏳ This may take 15-30 seconds per URL...")
            
            result = creator.setup_project(
                user_id=user_id,
                project_name=project_name,
                urls=urls_list,
                whitepaper_path=None
            )
            
            current_step += 1
            progress_bar.progress(current_step / total_steps)
            st.info(f"✅ URLs processed: {result}")
        
        # ═══ STEP 2: Combine text content ═══
        combined_text_content = ""
        
        if additional_text and additional_text.strip():
            status_text.text("📝 Processing pasted content...")
            combined_text_content += additional_text + "\n\n"
            current_step += 1
            progress_bar.progress(current_step / total_steps)
        
        # ═══ STEP 3: Process uploaded files ═══
        if uploaded_files:
            for i, uploaded_file in enumerate(uploaded_files):
                status_text.text(f"📄 Processing file {i+1}/{len(uploaded_files)}: {uploaded_file.name}")
                content = extract_file_content(uploaded_file)
                if content:
                    combined_text_content += f"\n--- Content from {uploaded_file.name} ---\n\n{content}\n\n"
                current_step += 1
                progress_bar.progress(current_step / total_steps)
        
        # ═══ STEP 4: Create embeddings for text content ═══
        if combined_text_content.strip():
            from langchain_core.documents import Document
            
            text_doc = Document(page_content=combined_text_content)
            chunks = creator.text_splitter.split_documents([text_doc])
            
            status_text.text(f"🧠 Creating embeddings for {len(chunks)} chunks...")
            status_text.caption(f"⏳ Estimated time: {len(chunks) * 0.04:.0f}-{len(chunks) * 0.08:.0f} seconds...")
            
            if chunks:
                persist_dir = creator._get_project_path(user_id, project_name)
                
                if user_id in creator.vector_store:
                    # Add to existing store with batching
                    batch_size = 100
                    for i in range(0, len(chunks), batch_size):
                        batch = chunks[i:i+batch_size]
                        creator.vector_store[user_id].add_documents(batch)
                        progress_current = min(i + batch_size, len(chunks))
                        status_text.text(f"🧠 Embedding progress: {progress_current}/{len(chunks)} chunks")
                    
                    st.info(f"✅ Added {len(chunks)} chunks to existing project")
                else:
                    # Create new store (only if no URLs were processed)
                    if not urls_list:
                        from langchain_community.vectorstores import Chroma
                        
                        store = Chroma.from_documents(
                            documents=chunks,
                            embedding=creator.embeddings,
                            collection_name=f"{user_id}_{project_name.lower()}",
                            persist_directory=persist_dir
                        )
                        creator.vector_store[user_id] = store
                        creator.project_name[user_id] = project_name
                        st.info(f"✅ Created new project with {len(chunks)} chunks")
        
        # ═══ COMPLETE ═══
        progress_bar.progress(1.0)
        status_text.text("✅ Processing complete!")
        time.sleep(0.5)
        progress_bar.empty()
        status_text.empty()
        
        # Return document count
        if user_id in creator.vector_store:
            return len(creator.vector_store[user_id].get()['ids'])
        return 0
        
    except Exception as e:
        st.error(f"Error setting up project: {str(e)}")
        import traceback
        st.error(traceback.format_exc())
        return 0


# ============================================
# MAIN APP
# ============================================

def main():
    st.set_page_config(
        page_title="AI Content Creator",
        page_icon="🧵",
        layout="wide"
    )
    
    # Initialize session state for email
    if 'email_provided' not in st.session_state:
        st.session_state.email_provided = False
        st.session_state.user_email = None
    
    # Show email gate if not authenticated
    if not st.session_state.email_provided:
        show_email_gate()
        st.stop()
    
    # Initialize auth system
    auth_system = UserAuthSystem(DATABASE_URL)
    
    # User is authenticated - show main app
    st.title("🧵 AI Twitter Content Creator")
    st.markdown("Create engaging, research-backed Twitter content for crypto projects")
    
    # Initialize session state
    if 'creator' not in st.session_state:
        try:
            st.session_state.creator = MultiUserContentCreator(GROQ_API_KEY)
        except Exception as e:
            st.error(f"Error initializing content creator: {str(e)}")
            st.stop()
    
    if 'data_uploaded' not in st.session_state:
        st.session_state.data_uploaded = False
    if 'project_name' not in st.session_state:
        st.session_state.project_name = ""
    if 'user_id' not in st.session_state:
        st.session_state.user_id = st.session_state.user_email.split('@')[0]
    
    # AUTO-REGISTER USER (FIXED)
    if st.session_state.user_id not in st.session_state.creator.user_profiles:
        st.session_state.creator.register_user(
            st.session_state.user_id,
            default_template="educational",
            brand_voice=""
        )
    
    # Sidebar for data upload and settings
    with st.sidebar:
        # User info at top
        st.markdown(f"### 👤 {st.session_state.user_email}")
        if st.button("🚪 Logout", use_container_width=True):
            st.session_state.email_provided = False
            st.session_state.user_email = None
            st.session_state.data_uploaded = False
            st.rerun()
        
        st.divider()
        st.header("📁 Project Setup")
        
        # User settings
        with st.expander("⚙️ User Settings", expanded=not st.session_state.data_uploaded):
            user_id = st.text_input(
                "User ID",
                value=st.session_state.user_id,
                help="Unique identifier for your account"
            )
            
            default_template = st.selectbox(
                "Default Template",
                list(TWEET_TEMPLATES.keys()),
                help="Your preferred content style"
            )
            
            brand_voice = st.text_area(
                "Brand Voice (optional)",
                placeholder="e.g., Technical expert, conversational but precise",
                help="Describe your desired tone and style"
            )
            
            if st.button("💾 Save User Settings"):
                st.session_state.user_id = user_id
                st.session_state.creator.register_user(
                    user_id,
                    default_template=default_template,
                    brand_voice=brand_voice if brand_voice else ""
                )
                st.success("✅ User settings saved!")
        
        st.divider()
        
        # Project data upload
        project_name = st.text_input(
            "Project Name",
            value=st.session_state.project_name,
            placeholder="e.g., bitcoin, ethereum, myproject"
        )
        
        # Load existing project option
        if project_name:
            projects = st.session_state.creator.list_user_projects(st.session_state.user_id)
            project_exists = any(p['project_name'].lower() == project_name.lower() for p in projects)
            
            if project_exists:
                st.info(f"📂 Project '{project_name}' exists")
                if st.button("📥 Load Existing Project"):
                    with st.spinner("Loading project..."):
                        success = st.session_state.creator.load_project(
                            st.session_state.user_id,
                            project_name
                        )
                        if success:
                            st.session_state.data_uploaded = True
                            st.session_state.project_name = project_name
                            st.success("✅ Project loaded!")
                            st.rerun()
        
        # URLs
        st.subheader("🔗 Website URLs")
        urls_text = st.text_area(
            "Add URLs (one per line)",
            height=100,
            placeholder="https://docs.yourproject.io\nhttps://blog.yourproject.io"
        )
        
        # File Upload
        st.subheader("📄 Upload Files")
        uploaded_files = st.file_uploader(
            "Choose files",
            accept_multiple_files=True,
            type=['txt', 'pdf', 'docx', 'md']
        )
        
        if uploaded_files:
            st.success(f"✅ {len(uploaded_files)} files ready")
            for file in uploaded_files:
                st.write(f"• {file.name}")
        
        # Additional Text
        st.subheader("📝 Paste Content")
        additional_text = st.text_area(
            "Paste any additional content",
            height=100,
            placeholder="Paste documentation, whitepapers, or any relevant text..."
        )
        
        # Upload Button
        if st.button("🚀 Process Data", type="primary"):
            if not project_name.strip():
                st.error("Please enter a project name!")
            else:
                urls_list = [url.strip() for url in urls_text.split('\n') if url.strip()]
                
                if not urls_list and not uploaded_files and not additional_text.strip():
                    st.error("Please add at least one data source!")
                else:
                    with st.spinner("Processing data..."):
                        try:
                            doc_count = process_project_data(
                                st.session_state.creator,
                                st.session_state.user_id,
                                project_name,
                                urls_list,
                                uploaded_files,
                                additional_text
                            )
                            
                            if doc_count > 0:
                                st.session_state.data_uploaded = True
                                st.session_state.project_name = project_name
                                st.success(f"✅ Data processed! {doc_count} chunks in knowledge base")
                            else:
                                st.error("❌ No data could be processed")
                        except Exception as e:
                            st.error(f"❌ Processing failed: {str(e)}")
        
        # Show existing projects
        if st.session_state.user_id:
            st.divider()
            projects = st.session_state.creator.list_user_projects(st.session_state.user_id)
            if projects:
                st.subheader("📚 Your Projects")
                for project in projects:
                    st.write(f"• {project['project_name']} ({project['urls_count']} sources)")
        
        st.divider()
        st.caption("💡 Free during beta • CPU optimized")
    
    # Main content area
    if not st.session_state.data_uploaded:
        st.info("👈 Please set up your project first using the sidebar")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("""
            ### 📝 Templates Available:
            - **Educational**: Deep dives & teaching
            - **Promotional**: Product launches
            - **Thread**: Multi-tweet stories
            - **Engagement**: Questions & discussions
            - **Casual**: Quick observations
            - **Storytelling**: Narratives & lessons
            """)
        
        with col2:
            st.markdown("""
            ### 🎯 How it works:
            1. Set your user preferences
            2. Add project name
            3. Upload documentation (URLs, files, text)
            4. Process the data (CPU optimized!)
            5. Create content!
            """)
        
        with col3:
            st.markdown("""
            ### 📄 Supported formats:
            - PDF files
            - Word documents (.docx)
            - Markdown files (.md)
            - Text files (.txt)
            - Website URLs
            - Direct text paste
            """)
    
    else:
        # Content creation interface
        st.header(f"✨ Create Content for {st.session_state.project_name}")
        
        # Show available templates
        with st.expander("📋 View Available Templates"):
            for key, template in TWEET_TEMPLATES.items():
                st.markdown(f"**{template['name']}** - {template['description']}")
                st.caption(f"Best for: {template['best_for']}")
                st.divider()
        
        # Content creation form
        topic = st.text_area(
            "What topic do you want to write about?",
            height=100,
            placeholder="e.g., 'How our ZK proof system achieves sub-second verification'"
        )
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            template = st.selectbox(
                "Content Template",
                list(TWEET_TEMPLATES.keys()),
                help="Choose the style of content"
            )
        
        with col2:
            length = st.selectbox(
                "Length",
                ["short", "medium", "long"],
                index=1,
                help="Short: 400-600 | Medium: 600-850 | Long: 900-1500 chars"
            )
        
        with col3:
            debug_mode = st.checkbox(
                "Debug Mode",
                help="Show research content being used"
            )
        
        if st.button("🎯 Generate Content", type="primary"):
            if not topic.strip():
                st.error("Please enter a topic!")
            else:
                # Track usage
                auth_system.increment_usage(st.session_state.user_email)
                
                with st.spinner("Creating content..."):
                    try:
                        content = st.session_state.creator.create_twitter_content(
                            user_id=st.session_state.user_id,
                            topic=topic,
                            template=template,
                            length=length,
                            debug=debug_mode
                        )
                        
                        if content and not content.startswith("Error:"):
                            st.success("✅ Content generated!")
                            
                            # Show stats
                            char_count = len(content)
                            if template == "thread":
                                tweet_count = content.count("---TWEET BREAK---") + 1
                                st.info(f"📊 {tweet_count} tweets | {char_count} total characters")
                            else:
                                st.info(f"📊 {char_count} characters")
                            
                            # Show content
                            st.text_area("Generated Content", content, height=400)
                            
                            # Download button
                            filename = f"{topic[:30].replace(' ', '_')}_{template}.txt"
                            st.download_button(
                                label="📥 Download Content",
                                data=content,
                                file_name=filename,
                                mime="text/plain"
                            )
                            
                            # Tweak option
                            st.divider()
                            st.subheader("✏️ Tweak Content")
                            feedback = st.text_area(
                                "What would you like to change?",
                                placeholder="e.g., 'Make it more technical' or 'Add more specific examples'"
                            )
                            
                            if st.button("🔄 Regenerate with Feedback"):
                                if feedback:
                                    with st.spinner("Tweaking content..."):
                                        tweaked = st.session_state.creator.tweak_content(
                                            st.session_state.user_id,
                                            content,
                                            feedback
                                        )
                                        st.text_area("Tweaked Content", tweaked, height=400)
                                        st.download_button(
                                            label="📥 Download Tweaked Content",
                                            data=tweaked,
                                            file_name=f"{filename.replace('.txt', '_tweaked.txt')}",
                                            mime="text/plain"
                                        )
                        else:
                            st.error("❌ Content generation failed")
                            if content:
                                st.error(content)
                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")
                        import traceback
                        st.error(traceback.format_exc())
        
        st.divider()
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔄 Start New Project"):
                st.session_state.data_uploaded = False
                st.session_state.project_name = ""
                st.rerun()
        with col2:
            if st.button("🗑️ Delete Current Project"):
                if st.session_state.creator.delete_project(
                    st.session_state.user_id,
                    st.session_state.project_name
                ):
                    st.session_state.data_uploaded = False
                    st.session_state.project_name = ""
                    st.success("Project deleted!")
                    st.rerun()


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        st.error(f"Application error: {str(e)}")
        st.error("Please refresh the page and try again.")